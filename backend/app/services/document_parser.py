"""Text extraction for original knowledge files."""

import csv
import io
import re


class DocumentParseError(ValueError):
    """Raised when an uploaded document cannot be parsed safely."""


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError

        document = Document(io.BytesIO(data))
        paragraphs: list[str] = []
        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text.strip()
            if not paragraph_text:
                continue
            style_name = (paragraph.style.name or "").lower()
            heading_match = re.match(r"heading\s*(\d+)", style_name)
            if heading_match:
                level = min(int(heading_match.group(1)), 6)
                paragraph_text = f"{'#' * level} {paragraph_text}"
            paragraphs.append(paragraph_text)
        for table in document.tables:
            for row in table.rows:
                values = [
                    re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells
                ]
                if any(values):
                    paragraphs.append(" | ".join(values))
        return "\n".join(text for text in paragraphs if text.strip())
    except (PackageNotFoundError, ValueError, KeyError) as exc:
        raise DocumentParseError("Invalid DOCX file") from exc


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentParseError("PDF parser is not installed") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join(
            f"[[PAGE:{page_number}]]\n{page.extract_text() or ''}"
            for page_number, page in enumerate(reader.pages, start=1)
        )
    except Exception as exc:
        raise DocumentParseError("Invalid or encrypted PDF file") from exc


def extract_file_text(filename: str, data: bytes) -> str:
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    try:
        if extension in {"txt", "md"}:
            return data.decode("utf-8-sig")
        if extension == "csv":
            decoded = data.decode("utf-8-sig")
            rows = csv.reader(io.StringIO(decoded))
            return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)
    except UnicodeDecodeError as exc:
        raise DocumentParseError("Text file must use UTF-8 encoding") from exc
    if extension == "docx":
        return _extract_docx(data)
    if extension == "pdf":
        return _extract_pdf(data)
    raise DocumentParseError("Supported file types: PDF, DOCX, TXT, MD and CSV")
