"""Generate governed, type-specific legal drafts as polished DOCX or PDF files."""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any

from app.services.legal_documents.schemas import DOCUMENT_SCHEMAS
from app.services.legal_documents.templates import build_document_draft
from app.services.legal_documents.validators import validate_document_fields


SUPPORTED_TEMPLATES = {
    document_type: schema["label"] for document_type, schema in DOCUMENT_SCHEMAS.items()
}


def _set_docx_font(run, *, name: str = "Arial", size: int | None = None, bold: bool | None = None, color: str | None = None, italic: bool | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_table_geometry(table, widths: list[int], *, indent: int = 120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    total_width = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(index, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
            margins = tc_pr.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def _style_docx_table(table, widths: list[int], *, header: bool = True) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Pt

    _set_table_geometry(table, widths)
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if header and row_index == 0:
                _set_cell_shading(cell, "F2F4F7")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    _set_docx_font(run, size=9, bold=header and row_index == 0)


def _docx_bytes(draft: dict[str, Any]) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color, before, after in (
        ("Heading 1", 13, "2E74B5", 12, 6),
        ("Heading 2", 12, "1F4D78", 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    _set_docx_font(header.add_run("AI WORKFORCE  |  LEGAL DRAFT"), size=8, bold=True, color="667085")
    _set_docx_font(header.add_run("                                      CONFIDENTIAL"), size=8, color="98A2B3")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_docx_font(footer.add_run("Bản nháp do AI tạo - Chỉ có hiệu lực sau khi Legal phê duyệt  |  Trang "), size=8, color="667085")
    page_field = OxmlElement("w:fldSimple")
    page_field.set(qn("w:instr"), "PAGE")
    footer._p.append(page_field)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(4)
    _set_docx_font(kicker.add_run("DỰ THẢO - CẦN PHÊ DUYỆT PHÁP LÝ"), size=9, bold=True, color="9B1C1C")
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    _set_docx_font(title.add_run(draft["title"]), size=22, bold=True, color="0B2545")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    _set_docx_font(subtitle.add_run(draft["subtitle"]), size=12, color="475467")

    metadata = document.add_table(rows=0, cols=2)
    for label, value in [("Bên A", draft["party_a"]), ("Bên B", draft["party_b"]), *draft["metadata"]]:
        cells = metadata.add_row().cells
        cells[0].text = label
        cells[1].text = value
    _style_docx_table(metadata, [2700, 6660], header=False)
    for row in metadata.rows:
        _set_cell_shading(row.cells[0], "E8EEF5")
        for run in row.cells[0].paragraphs[0].runs:
            _set_docx_font(run, size=9, bold=True, color="1F4D78")

    intro = document.add_paragraph()
    intro.paragraph_format.space_before = Pt(12)
    intro.paragraph_format.space_after = Pt(8)
    _set_docx_font(intro.add_run(f"Các Bên gồm {draft['party_a']} và {draft['party_b']} thống nhất ký kết văn bản này với các điều khoản sau:"), size=11)

    for index, item in enumerate(draft["sections"], 1):
        heading = document.add_heading(f"{index}. {item['title'].upper()}", level=1)
        heading.paragraph_format.keep_with_next = True
        for paragraph_text in item.get("paragraphs", []):
            paragraph = document.add_paragraph(paragraph_text)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for bullet in item.get("bullets", []):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.1
            _set_docx_font(paragraph.add_run(bullet), size=11)
        if table_data := item.get("table"):
            headers = table_data["headers"]
            rows = table_data["rows"]
            table = document.add_table(rows=1, cols=len(headers))
            for column, value in enumerate(headers):
                table.rows[0].cells[column].text = value
            for row_values in rows:
                cells = table.add_row().cells
                padded = list(row_values) + [""] * (len(headers) - len(row_values))
                for column, value in enumerate(padded[: len(headers)]):
                    cells[column].text = str(value).strip()
            widths = [1800, 4680, 2880] if len(headers) == 3 else [9360 // len(headers)] * len(headers)
            _style_docx_table(table, widths)

    if draft["warnings"]:
        heading = document.add_heading("CẢNH BÁO RÀ SOÁT PHÁP LÝ", level=1)
        heading.paragraph_format.keep_with_next = True
        for warning in draft["warnings"]:
            box = document.add_table(rows=1, cols=1)
            cell = box.cell(0, 0)
            cell.text = ""
            _set_cell_shading(cell, "FDECEC")
            title_paragraph = cell.paragraphs[0]
            _set_docx_font(title_paragraph.add_run(f"{warning['severity']} - {warning['title']}"), size=10, bold=True, color="9B1C1C")
            detail = cell.add_paragraph(warning["message"])
            recommendation = cell.add_paragraph(f"Khuyến nghị: {warning['recommendation']}")
            for paragraph in (detail, recommendation):
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    _set_docx_font(run, size=9)
            _set_table_geometry(box, [9360])

    document.add_heading("CHỮ KÝ", level=1)
    signature = document.add_table(rows=2, cols=2)
    signature.cell(0, 0).text = f"ĐẠI DIỆN BÊN A\n{draft['party_a']}"
    signature.cell(0, 1).text = f"ĐẠI DIỆN BÊN B\n{draft['party_b']}"
    signature.cell(1, 0).text = "\n\n\nHọ tên: ____________________\nChức danh: _________________\nNgày: ______________________"
    signature.cell(1, 1).text = "\n\n\nHọ tên: ____________________\nChức danh: _________________\nNgày: ______________________"
    _style_docx_table(signature, [4680, 4680], header=False)
    for cell in signature.rows[0].cells:
        _set_cell_shading(cell, "F2F4F7")
        for run in cell.paragraphs[0].runs:
            _set_docx_font(run, size=10, bold=True, color="1F4D78")

    notice = document.add_paragraph()
    notice.paragraph_format.space_before = Pt(12)
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_docx_font(notice.add_run("LƯU Ý: Đây là bản nháp do AI hỗ trợ tạo. Văn bản phải được Legal kiểm tra, phê duyệt và hoàn thiện thông tin pháp lý trước khi ký."), size=8, bold=True, color="9B1C1C")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _register_pdf_fonts() -> tuple[str, str]:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = [
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        (r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\arialbd.ttf"),
    ]
    for regular_path, bold_path in candidates:
        if Path(regular_path).exists() and Path(bold_path).exists():
            pdfmetrics.registerFont(TTFont("LegalSans", regular_path))
            pdfmetrics.registerFont(TTFont("LegalSans-Bold", bold_path))
            return "LegalSans", "LegalSans-Bold"
    return "Helvetica", "Helvetica-Bold"


def _pdf_bytes(draft: dict[str, Any]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    regular_font, bold_font = _register_pdf_fonts()
    output = io.BytesIO()
    pdf = SimpleDocTemplate(output, pagesize=LETTER, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()
    body = ParagraphStyle("LegalBody", parent=styles["BodyText"], fontName=regular_font, fontSize=10, leading=13, spaceAfter=6, alignment=TA_JUSTIFY, textColor=colors.HexColor("#1D2939"))
    heading = ParagraphStyle("LegalHeading", parent=styles["Heading1"], fontName=bold_font, fontSize=12, leading=15, spaceBefore=12, spaceAfter=6, textColor=colors.HexColor("#2E74B5"), keepWithNext=True)
    small = ParagraphStyle("LegalSmall", parent=body, fontSize=8, leading=10, textColor=colors.HexColor("#667085"))
    warning_style = ParagraphStyle("LegalWarning", parent=body, fontSize=9, leading=12, textColor=colors.HexColor("#9B1C1C"))
    story: list[Any] = []
    story.append(Paragraph("DỰ THẢO - CẦN PHÊ DUYỆT PHÁP LÝ", ParagraphStyle("Kicker", parent=small, fontName=bold_font, textColor=colors.HexColor("#9B1C1C"), spaceAfter=5)))
    story.append(Paragraph(draft["title"], ParagraphStyle("Title", parent=styles["Title"], fontName=bold_font, fontSize=20, leading=24, alignment=0, textColor=colors.HexColor("#0B2545"), spaceAfter=4)))
    story.append(Paragraph(draft["subtitle"], ParagraphStyle("Subtitle", parent=body, fontSize=11, textColor=colors.HexColor("#475467"), spaceAfter=14)))
    metadata_rows = [[Paragraph(f"<b>{label}</b>", body), Paragraph(value, body)] for label, value in [("Bên A", draft["party_a"]), ("Bên B", draft["party_b"]), *draft["metadata"]]]
    metadata = Table(metadata_rows, colWidths=[1.75 * inch, 4.75 * inch], repeatRows=0)
    metadata.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D0D5DD")), ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E8EEF5")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 7), ("RIGHTPADDING", (0, 0), (-1, -1), 7), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
    story.extend([metadata, Spacer(1, 12), Paragraph(f"Các Bên gồm {draft['party_a']} và {draft['party_b']} thống nhất ký kết văn bản này với các điều khoản sau:", body)])
    for index, item in enumerate(draft["sections"], 1):
        story.append(Paragraph(f"{index}. {item['title'].upper()}", heading))
        for text in item.get("paragraphs", []):
            story.append(Paragraph(text.replace("\n", "<br/>"), body))
        for bullet in item.get("bullets", []):
            story.append(Paragraph(f"• {bullet}", ParagraphStyle("LegalBullet", parent=body, leftIndent=18, firstLineIndent=-10, spaceAfter=4)))
        if table_data := item.get("table"):
            headers = [Paragraph(f"<b>{value}</b>", small) for value in table_data["headers"]]
            rows = [[Paragraph(str(value), small) for value in row] for row in table_data["rows"]]
            table = Table([headers, *rows], colWidths=[1.25 * inch, 3.25 * inch, 2 * inch], repeatRows=1)
            table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D0D5DD")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6), ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5)]))
            story.append(table)
    if draft["warnings"]:
        story.append(Paragraph("CẢNH BÁO RÀ SOÁT PHÁP LÝ", heading))
        for warning in draft["warnings"]:
            content = Paragraph(f"<b>{warning['severity']} - {warning['title']}</b><br/>{warning['message']}<br/><b>Khuyến nghị:</b> {warning['recommendation']}", warning_style)
            box = Table([[content]], colWidths=[6.5 * inch])
            box.setStyle(TableStyle([("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#ECA5A5")), ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FDECEC")), ("LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 8), ("TOPPADDING", (0, 0), (-1, -1), 7), ("BOTTOMPADDING", (0, 0), (-1, -1), 7)]))
            story.extend([box, Spacer(1, 6)])
    story.append(Paragraph("CHỮ KÝ", heading))
    signature_data = [[Paragraph(f"<b>ĐẠI DIỆN BÊN A</b><br/>{draft['party_a']}", body), Paragraph(f"<b>ĐẠI DIỆN BÊN B</b><br/>{draft['party_b']}", body)], [Paragraph("<br/><br/><br/>Họ tên: ____________________<br/>Chức danh: _________________<br/>Ngày: ______________________", body), Paragraph("<br/><br/><br/>Họ tên: ____________________<br/>Chức danh: _________________<br/>Ngày: ______________________", body)]]
    signature = Table(signature_data, colWidths=[3.25 * inch, 3.25 * inch])
    signature.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D0D5DD")), ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")), ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 8), ("RIGHTPADDING", (0, 0), (-1, -1), 8), ("TOPPADDING", (0, 0), (-1, -1), 7), ("BOTTOMPADDING", (0, 0), (-1, -1), 7)]))
    story.extend([signature, Spacer(1, 12), Paragraph("LƯU Ý: Đây là bản nháp do AI hỗ trợ tạo. Văn bản phải được Legal kiểm tra, phê duyệt và hoàn thiện thông tin pháp lý trước khi ký.", ParagraphStyle("Notice", parent=small, fontName=bold_font, textColor=colors.HexColor("#9B1C1C"), alignment=TA_CENTER))])

    def page_furniture(canvas, doc):
        canvas.saveState()
        canvas.setFont(regular_font, 8)
        canvas.setFillColor(colors.HexColor("#667085"))
        canvas.drawString(inch, LETTER[1] - 0.55 * inch, "AI WORKFORCE  |  LEGAL DRAFT")
        canvas.drawRightString(LETTER[0] - inch, 0.55 * inch, f"Trang {doc.page}")
        canvas.restoreState()

    pdf.build(story, onFirstPage=page_furniture, onLaterPages=page_furniture)
    return output.getvalue()


def render_contract_text(document_type: str, fields: dict[str, Any]) -> str:
    """Return a readable text representation for compatibility and tests."""
    validation = validate_document_fields(document_type, fields)
    if not validation["valid"]:
        missing = ", ".join(item["label"] for item in validation["missing_fields"])
        raise ValueError(f"Thiếu thông tin bắt buộc: {missing}")
    draft = build_document_draft(document_type, validation["normalized_fields"], validation["warnings"])
    lines = [draft["title"], draft["subtitle"]]
    for index, section in enumerate(draft["sections"], 1):
        lines.extend([f"{index}. {section['title']}", *section.get("paragraphs", []), *section.get("bullets", [])])
    return "\n\n".join(lines)


def generate_legal_document(document_type: str, output_format: str, fields: dict[str, Any]) -> tuple[bytes, str, str]:
    normalized_type = document_type.upper()
    normalized_format = output_format.lower()
    validation = validate_document_fields(normalized_type, fields)
    if not validation["valid"]:
        missing = ", ".join(item["label"] for item in validation["missing_fields"])
        raise ValueError(f"Thiếu thông tin bắt buộc: {missing}")
    draft = build_document_draft(normalized_type, validation["normalized_fields"], validation["warnings"])
    safe_name = normalized_type.lower()
    if normalized_format == "docx":
        return _docx_bytes(draft), f"{safe_name}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if normalized_format == "pdf":
        return _pdf_bytes(draft), f"{safe_name}.pdf", "application/pdf"
    raise ValueError("Định dạng đầu ra phải là DOCX hoặc PDF")
